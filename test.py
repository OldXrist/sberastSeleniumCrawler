import os
import docx

def edit_application(procedure_number, quantity):
    path_to_doc = os.getcwd().replace('\\', '/') + '/docs-aktiv/application/Заявка Актив.docx'
    doc = docx.Document(path_to_doc)
    total = '0,01'

    if quantity != '':
        total = str(0.01 * float(quantity))

    doc.paragraphs[2].runs[2].text = procedure_number
    doc.paragraphs[9].runs[1].text = procedure_number
    doc.paragraphs[9].runs[3].text = total
    doc.paragraphs[11].runs[1].text = quantity.split('.')[0]

    doc.save(os.getcwd().replace('\\', '/') + '/docs-aktiv/application/Заявка Актив-test.docx')

edit_application('123', '4')