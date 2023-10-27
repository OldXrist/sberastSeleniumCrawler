from selenium import webdriver
from selenium.webdriver import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.action_chains import ActionChains
import time
import os
import docx
import pyautogui as pygui


def edit_application(procedure_number, quantity):
    path_to_doc = os.getcwd().replace('\\', '/') + '/docs-aktiv/application/Заявка Актив.docx'
    doc = docx.Document(path_to_doc)
    total = '0,01'

    if quantity != '':
        total = str(0.01 * float(quantity))

    doc.paragraphs[2].runs[2].text = procedure_number
    doc.paragraphs[9].runs[1].text = procedure_number
    doc.paragraphs[9].runs[3].text = total
    doc.paragraphs[11].runs[1].text = quantity.split('.')[0]

    doc.save(os.getcwd().replace('\\', '/') + '/docs-aktiv/application/Заявка Актив.docx')


def click_await(driver, selector, index):
    while True:
        if driver.find_elements(By.CSS_SELECTOR, '#ajax-background') and driver.find_elements(By.CSS_SELECTOR, '#ajax-background')[0].is_displayed():
            time.sleep(1)
        else:
            if driver.find_elements(By.CSS_SELECTOR, selector) and driver.find_elements(By.CSS_SELECTOR, selector)[index].is_displayed():
                driver.find_elements(By.CSS_SELECTOR, selector)[index].click()
                break
            elif driver.find_elements(By.CSS_SELECTOR, selector) and not driver.find_elements(By.CSS_SELECTOR, selector)[index].is_displayed():
                break
            else:
                time.sleep(1)


def send_keys_by_css(driver, selector, index, keys):
    while True:
        if driver.find_elements(By.CSS_SELECTOR, '#ajax-background') and driver.find_elements(By.CSS_SELECTOR, '#ajax-background')[0].is_displayed():
            time.sleep(1)
        else:
            if driver.find_elements(By.CSS_SELECTOR, selector) and driver.find_elements(By.CSS_SELECTOR, selector)[index].is_displayed():
                driver.find_elements(By.CSS_SELECTOR, selector)[index].send_keys(keys)
                break
            else:
                time.sleep(1)


def main(driver):
    while True:
        if driver.find_elements(By.CSS_SELECTOR, '#mainContent_btnSignInByERUZ'):
            pygui.press('enter')
            break
        else:
            time.sleep(1)

    time.sleep(2)

    click_await(driver, '#mainContent_btnSignInByERUZ', 0)
    click_await(driver, '#btnEnter', 0)
    click_await(driver, '#searchClear', 0)
    driver.execute_script("document.getElementById('ajax-background').remove()")
    send_keys_by_css(driver, '#searchInput', 0, 'Запрос котировок в электронной форме')
    click_await(driver, '.simple-button.orange-background', 0)
    click_await(driver, '#expandAdditionalFilters', 0)
    time.sleep(1)
    click_await(driver, '#customerOpenButton', 0)
    time.sleep(1)
    click_await(driver, '#textSearchLongDict', 0)
    send_keys_by_css(driver, '#textSearchLongDict', 0, '7708701670')
    driver.execute_script('applyOrgFilter()')
    driver.execute_script('applyMainFilters()')

    time.sleep(2)

    while True:
        if driver.find_elements(By.CSS_SELECTOR, ".purch-reestr-tbl-div"):
            lotNum = driver.find_elements(By.CSS_SELECTOR, ".es-el-code-term")[0].get_attribute("innerText")
            with open('lot_numbers.txt', 'r') as r:
                if lotNum in r.read():
                    driver.execute_script("document.getElementsByClassName('purch-reestr-tbl-div')[0].remove()")
                else:
                    with open('lot_numbers.txt', 'a') as f:
                        f.write(lotNum)
                        if driver.find_elements(By.CSS_SELECTOR, ".link-button"):
                            driver.find_elements(By.CSS_SELECTOR, ".link-button")[0].click()
                        break
        else:
            driver.execute_script("handleMainSearch()")
            time.sleep(3)

    while True:
        if driver.find_elements(By.CSS_SELECTOR, '#ctl00_ctl00_phWorkZone_BackToParent'):
            break
        else:
            time.sleep(1)

    pygui.press('tab')
    time.sleep(1)
    pygui.press('enter')

    trade_num = ''

    if driver.find_elements(By.CSS_SELECTOR, '.dt'):
        trade_num = driver.find_elements(By.CSS_SELECTOR, '.dt')[2].find_elements(By.CSS_SELECTOR, 'td')[2].get_attribute('innerText')

    if driver.find_elements(By.CSS_SELECTOR, '.hiddenContent'):
        driver.execute_script('document.getElementsByClassName("hiddenContent")[0].style.display = "block"')

    quantity = driver.find_elements(By.CSS_SELECTOR, 'span')[20].get_attribute('innerText')

    edit_application(trade_num, quantity)

    click_await(driver, '#OpenDictbxAccFake', 0)
    click_await(driver, '#OpenDictbxAcc', 0)

    driver.switch_to.frame(driver.find_elements(By.CSS_SELECTOR, '#spravIframe')[0])
    driver.find_elements(By.CSS_SELECTOR, '#XMLContainer')[0].find_elements(By.CSS_SELECTOR, 'a')[0].click()

    driver.switch_to.default_content()

    add_doc_buttons = driver.find_elements(By.CSS_SELECTOR, '.btnElastic')

    if add_doc_buttons[0].get_attribute('value') == 'Прикрепить':
        add_doc_buttons[0].click()
        driver.switch_to.frame(driver.find_elements(By.CSS_SELECTOR, '#spravIframe')[0])
        send_keys_by_css(driver, '#ctl00_phDataZone_Upload', 0, os.getcwd().replace('\\', '/') + '/docs-aktiv/application/Заявка Актив.docx')
        click_await(driver, '#ctl00_phDataZone_btnDoUpload', 0)

    driver.switch_to.default_content()

    k = 1

    if add_doc_buttons[1].get_attribute('value') == 'Прикрепить':
        add_doc_buttons[1].click()
        driver.switch_to.frame(driver.find_elements(By.CSS_SELECTOR, '#spravIframe')[0])
        send_keys_by_css(driver, '#ctl00_phDataZone_Upload', 0, os.getcwd().replace('\\', '/') + '/docs-aktiv/docs-1/Выписка ЕГРЮЛ.pdf')
        click_await(driver, '#ctl00_phDataZone_btnDoUpload', 0)
        driver.switch_to.default_content()

    for filename in os.listdir(os.getcwd().replace('\\', '/') + '/docs/docs-2'):
        driver.find_elements(By.CSS_SELECTOR, '#conformityProductDoc')[0].find_elements(By.CSS_SELECTOR, 'a')[k].click()
        driver.switch_to.frame(driver.find_elements(By.CSS_SELECTOR, '#spravIframe')[0])
        send_keys_by_css(driver, '#ctl00_phDataZone_Upload', 0, os.getcwd().replace('\\', '/') + '/docs-aktiv/docs-2/' + filename)
        click_await(driver, '#ctl00_phDataZone_btnDoUpload', 0)
        driver.switch_to.default_content()
        k += 1

    if add_doc_buttons[4].get_attribute('value') == 'Прикрепить':
        add_doc_buttons[4].click()
        driver.switch_to.frame(driver.find_elements(By.CSS_SELECTOR, '#spravIframe')[0])
        send_keys_by_css(driver, '#ctl00_phDataZone_Upload', 0, os.getcwd().replace('\\', '/') + '/docs-aktiv/deal_approval/Об одобрении крупнои сделки.pdf')
        click_await(driver, '#ctl00_phDataZone_btnDoUpload', 0)
        driver.switch_to.default_content()

    if add_doc_buttons[5].get_attribute('value') == 'Прикрепить':
        add_doc_buttons[5].click()
        driver.switch_to.frame(driver.find_elements(By.CSS_SELECTOR, '#spravIframe')[0])
        send_keys_by_css(driver, '#ctl00_phDataZone_Upload', 0, os.getcwd().replace('\\', '/') + '/docs-aktiv/declaration-1/Декларация.pdf')
        click_await(driver, '#ctl00_phDataZone_btnDoUpload', 0)
        driver.switch_to.default_content()

    if add_doc_buttons[6].get_attribute('value') == 'Загрузить':
        add_doc_buttons[6].click()
        driver.switch_to.frame(driver.find_elements(By.CSS_SELECTOR, '#spravIframe')[0])
        click_await(driver, '#ctl00_phDataZone_createTemplate', 0)
        driver.switch_to.default_content()

    if driver.find_elements(By.CSS_SELECTOR, '#ctl00_ctl00_phWorkZone_phDocumentZone_nbtPurchaseRequest_bankaccountTd'):
        driver.execute_script("document.getElementById('ctl00_ctl00_phWorkZone_phDocumentZone_nbtPurchaseRequest_bankaccountTd').getElementsByTagName('input')[0].click()")

    total = '0.01'
    if quantity != '':
        total = str(float(quantity) * 0.01)

    send_keys_by_css(driver, '#ContractAmount', 0, total)

    time.sleep(10000)


if __name__ == '__main__':
    options = Options()
    cwd = os.getcwd().replace('\\', '/') + "/UserData"
    options.add_argument("--user-data-dir=" + cwd)
    options.page_load_strategy = 'normal'
    chromeDriver = webdriver.Chrome(options=options)
    chromeDriver.get("https://www.sberbank-ast.ru/tradezone/Supplier/ESPurchaseList.aspx")
    main(chromeDriver)
